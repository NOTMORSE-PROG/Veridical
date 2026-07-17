"""V-006 reference-extraction tests.

The >=90% acceptance measurement runs against the owner's real reference
list with hand-verified labels in context/golden/demo_references.json —
both local-only (D-007), so CI skips that suite and covers the parser with
a synthetic well-formed APA fixture instead.
"""

import json
from pathlib import Path

import pytest

from app.config import get_settings
from app.ingest.patterns import load_patterns
from app.ingest.references import extract_references, parse_reference
from app.models.enums import CitationParseStatus
from tests.test_ingest_pdf import DEMO_PDF, PdfBuilder

GOLDEN = Path(__file__).resolve().parents[2] / "context" / "golden" / "demo_references.json"

# Synthetic, authored for the tests — generic shapes, not real papers.
WELL_FORMED = [
    (
        "Reyes, J. P., & Cruz, M. A. (2023). Assessing capstone readiness in state "
        "universities. Philippine Journal of Education, 12(3), 45–67. "
        "https://doi.org/10.1234/pje.2023.045",
        {"first": "Reyes, J. P.", "n": 2, "year": 2023, "doi": "10.1234/pje.2023.045"},
    ),
    (
        "Garcia, L. (2020). Understanding rubric design (2nd ed.). Academic Press. "
        "ISBN 978-0-12-345678-9",
        {"first": "Garcia, L.", "n": 1, "year": 2020, "isbn": "9780123456789"},
    ),
    (
        "Commission on Higher Education. (2019). CMO No. 15 Series 2019: Policies for "
        "graduate programs. CHED.",
        {"first": "Commission on Higher Education", "n": 1, "year": 2019, "doi": None},
    ),
    (
        "Tan, K., Lim, R. S., & Uy, D. (2024, March 5). Statistical errors in student "
        "theses. Journal of Academic Integrity, 8(1), 10–22. "
        "https://doi.org/10.5678/jai.2024.010",
        {"first": "Tan, K.", "n": 3, "year": 2024, "doi": "10.5678/jai.2024.010"},
    ),
    (
        "Lopez, A. B. (2021). Citation practices among undergraduates (arXiv:2101.00001). "
        "arXiv. https://arxiv.org/abs/2101.00001",
        {"first": "Lopez, A. B.", "n": 1, "year": 2021, "doi": None},
    ),
]


@pytest.mark.parametrize(("raw", "want"), WELL_FORMED, ids=[w[1]["first"] for w in WELL_FORMED])
def test_well_formed_apa_entries_parse(raw, want):
    d = parse_reference(raw, 0)
    assert d.parse_status == CitationParseStatus.parsed
    assert d.authors[0] == want["first"] and len(d.authors) == want["n"]
    assert d.year == want["year"]
    assert d.title and d.raw_text == raw
    if want.get("doi") is not None:
        assert d.doi == want["doi"]
    if want.get("isbn"):
        assert d.isbn == want["isbn"]


def test_arxiv_venue_split():
    d = parse_reference(WELL_FORMED[4][0], 0)
    assert d.title == "Citation practices among undergraduates (arXiv:2101.00001)"
    assert d.venue == "arXiv"


def test_malformed_entry_preserved_raw_and_flagged():
    raw = "some scribble that is not a citation at all"
    d = parse_reference(raw, 3)
    assert d.parse_status == CitationParseStatus.parse_failed
    assert d.raw_text == raw and d.order_index == 3


@pytest.mark.parametrize(
    "garbage",
    [
        "",
        "   ",
        "(((((((((",
        "​​​",
        "10.9999/orphan-doi-with-no-year",
        "Really long nonsense " * 500,
        "Åuthor, Ø. (notayear). Broken (2026",
        "\\x00\\x01 binary-ish ISBN: garbage ISBN 12",
    ],
)
def test_parser_never_throws(garbage):
    d = parse_reference(garbage, 0)
    assert d.raw_text == garbage  # whatever happens, the entry survives


def test_pdf_hanging_indent_segmentation(tmp_path):
    """Entries wrap 2-3 lines with continuations indented deeper — the
    section's leftmost x0 marks entry starts (the demo document's layout)."""
    b = PdfBuilder()
    b.new_page().line("CHAPTER 1 INTRODUCTION", bold=True).line("Body text about things.")
    b.new_page().line("REFERENCES", bold=True)
    entries = [
        [
            "Reyes, J. P., & Cruz, M. A. (2023). Assessing capstone readiness in",
            "state universities. Philippine Journal of Education, 12(3), 45–67.",
        ],
        ["Garcia, L. (2020). Understanding rubric design (2nd ed.). Academic", "Press."],
        [
            "Tan, K. (2024). Statistical errors in student theses. Journal of",
            "Academic Integrity, 8(1), 10–22.",
        ],
    ]
    for first, *rest in entries:
        b.line(first)
        for cont in rest:
            b.page.insert_text((100, b.y), cont, fontsize=11, fontname="helv")
            b.y += 18
    path = b.save(tmp_path / "refs.pdf")
    from app.ingest.pdf import extract_document

    settings = get_settings()
    result = extract_document(str(path), settings)
    drafts = extract_references(result, load_patterns(settings.ingest_patterns_file))
    assert [d.authors[0] for d in drafts] == ["Reyes, J. P.", "Garcia, L.", "Tan, K."]
    assert all(d.parse_status == CitationParseStatus.parsed for d in drafts)
    # Wrapped lines were rejoined into one raw entry.
    assert "Philippine Journal of Education" in drafts[0].raw_text


def test_docx_paragraph_segmentation(tmp_path):
    import docx as docx_lib

    from app.ingest.docx import extract_document

    d = docx_lib.Document()
    p = d.add_paragraph()
    p.add_run("REFERENCES").bold = True
    for raw, _ in WELL_FORMED[:3]:
        d.add_paragraph(raw)
    path = tmp_path / "refs.docx"
    d.save(str(path))
    settings = get_settings()
    result = extract_document(str(path), settings)
    drafts = extract_references(result, load_patterns(settings.ingest_patterns_file))
    assert len(drafts) == 3
    assert [d_.parse_status for d_ in drafts] == [CitationParseStatus.parsed] * 3


def test_document_without_references_yields_empty_not_error(tmp_path):
    b = PdfBuilder()
    b.new_page().line("CHAPTER 1 INTRODUCTION", bold=True).line("Just prose.")
    path = b.save(tmp_path / "norefs.pdf")
    from app.ingest.pdf import extract_document

    settings = get_settings()
    result = extract_document(str(path), settings)
    assert extract_references(result, load_patterns(settings.ingest_patterns_file)) == []


# --- the acceptance measurement on the real document (local-only) ------------


@pytest.mark.skipif(
    not (DEMO_PDF.exists() and GOLDEN.exists()),
    reason="demo PDF + golden labels are local-only (D-007)",
)
def test_demo_reference_list_accuracy_at_least_90_percent():
    from app.ingest.pdf import extract_document

    settings = get_settings()
    result = extract_document(str(DEMO_PDF), settings)
    drafts = extract_references(result, load_patterns(settings.ingest_patterns_file))
    expected = json.loads(GOLDEN.read_text(encoding="utf-8"))["expected"]
    assert len(drafts) == len(expected) == 17

    correct = 0
    for d, want in zip(drafts, expected, strict=True):
        if (
            (d.authors or [None])[0] == want["first_author"]
            and len(d.authors or []) == want["n_authors"]
            and d.year == want["year"]
            and d.title == want["title"]
            and d.doi == want["doi"]
            and d.url == want["url"]
        ):
            correct += 1
    accuracy = correct / len(expected)
    assert accuracy >= 0.9, f"reference accuracy {accuracy:.0%} below the F1.5 target"
