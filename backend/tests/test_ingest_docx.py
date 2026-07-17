"""V-005 DOCX ingestion tests.

No real capstone DOCX exists locally (the proposal is PDF-only), so the
PDF-equivalence acceptance criterion is exercised with a same-content
synthetic pair; a TIP-template DOCX should join the V-008 fixture corpus
when the owner obtains one (recorded gap).
"""

import io
from pathlib import Path

import docx as docx_lib
import pymupdf
import pytest

from app.config import get_settings
from app.errors import FileMalformedError
from app.ingest.docx import extract_document
from app.ingest.schemas import SectionNode
from app.ingest.service import select_extractor
from tests.test_ingest_pdf import PdfBuilder

# One logical document, produced in both formats for the equivalence test.
OUTLINE = [
    (1, "CHAPTER 1 INTRODUCTION"),
    (2, "1.1 Background of the Study"),
    (1, "CHAPTER 2 METHODS"),
    (2, "2.1 Design"),
    (3, "2.1.1 Sampling"),
    (1, "CHAPTER 3 RESULTS"),
]
BODY_TEXT = "Body prose that talks about the study at length."


def _png_bytes() -> bytes:
    pix = pymupdf.Pixmap(pymupdf.csRGB, pymupdf.IRect(0, 0, 8, 8))
    pix.clear_with(90)
    return pix.tobytes("png")


def _docx_with_styles(path: Path) -> Path:
    d = docx_lib.Document()
    for level, title in OUTLINE:
        d.add_heading(title, level=level)
        d.add_paragraph(BODY_TEXT)
    d.add_picture(io.BytesIO(_png_bytes()))
    table = d.add_table(rows=2, cols=2)
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            cell.text = f"cell {r}{c}"
    d.save(str(path))
    return path


def _docx_with_manual_bold(path: Path) -> Path:
    """The realistic student document: no Word styles, bold headings only."""
    d = docx_lib.Document()
    for _, title in OUTLINE:
        p = d.add_paragraph()
        p.add_run(title).bold = True
        d.add_paragraph(BODY_TEXT)
    d.save(str(path))
    return path


def _flatten(nodes: list[SectionNode]) -> list[tuple[int, str, str | None]]:
    out = []
    for n in nodes:
        out.append((n.level, n.title, n.numbering))
        out.extend(_flatten(n.children))
    return out


@pytest.fixture(scope="module")
def styled(tmp_path_factory):
    path = _docx_with_styles(tmp_path_factory.mktemp("docx") / "styled.docx")
    return extract_document(str(path), get_settings())


def test_word_heading_styles_build_the_tree(styled):
    tree = styled.section_tree
    assert tree.source == "styles"
    assert _flatten(tree.nodes) == [
        (1, "CHAPTER 1 INTRODUCTION", "1"),
        (2, "1.1 Background of the Study", "1.1"),
        (1, "CHAPTER 2 METHODS", "2"),
        (2, "2.1 Design", "2.1"),
        (3, "2.1.1 Sampling", "2.1.1"),
        (1, "CHAPTER 3 RESULTS", "3"),
    ]


def test_manual_bold_document_falls_back_to_heuristics(tmp_path):
    """Riskiest assumption of the ticket: most students never touch the
    style gallery — bold + numbering alone must recover the tree."""
    path = _docx_with_manual_bold(tmp_path / "manual.docx")
    tree = extract_document(str(path), get_settings()).section_tree
    assert tree.source == "heuristics"
    assert _flatten(tree.nodes) == _flatten(
        extract_document(
            str(_docx_with_styles(tmp_path / "styled.docx")), get_settings()
        ).section_tree.nodes
    )


def test_every_block_carries_a_paragraph_anchor(styled):
    assert styled.anchor_kind == "paragraph"
    assert styled.blocks
    assert all(b.paragraph is not None and b.page is None for b in styled.blocks)
    assert all(n.paragraph is not None for n in styled.section_tree.nodes)


def test_native_tables_extracted_with_rows_and_cells_intact(styled):
    assert len(styled.tables) == 1
    assert styled.tables[0].rows == [["cell 00", "cell 01"], ["cell 10", "cell 11"]]
    assert styled.tables[0].paragraph is not None


def test_images_inventoried_with_paragraph_anchor(styled):
    assert len(styled.images) == 1
    assert styled.images[0].paragraph is not None


def test_geometry_reports_declared_page_setup(styled):
    assert styled.geometry, "expected at least one Word section"
    g = styled.geometry[0]
    assert g.docx_section == 1 and g.page is None
    assert g.width == pytest.approx(612, abs=1)  # Letter, in points
    assert g.margins is not None and all(m > 0 for m in g.margins)


def test_docx_and_pdf_exports_yield_equivalent_trees(tmp_path):
    """Acceptance criterion: same document, two formats, one tree shape.
    Anchors differ by design (pages vs paragraphs) — structure must not."""
    from app.ingest.pdf import extract_document as extract_pdf

    docx_path = _docx_with_manual_bold(tmp_path / "same.docx")
    b = PdfBuilder()
    for _, title in OUTLINE:
        b.new_page().line(title, bold=True)
        b.line(BODY_TEXT)
    pdf_path = b.save(tmp_path / "same.pdf")

    docx_tree = extract_document(str(docx_path), get_settings()).section_tree
    pdf_tree = extract_pdf(str(pdf_path), get_settings()).section_tree
    assert _flatten(docx_tree.nodes) == _flatten(pdf_tree.nodes)


def test_toc_styled_paragraphs_do_not_become_headings(tmp_path):
    d = docx_lib.Document()
    # Word marks generated-TOC entries with "TOC N" styles ("toc 1" in
    # python-docx's style name casing).
    entry = d.add_paragraph("CHAPTER 1 INTRODUCTION 5")
    entry.style = d.styles["TOC 1"] if "TOC 1" in [s.name for s in d.styles] else entry.style
    d.add_heading("CHAPTER 1 INTRODUCTION", level=1)
    d.add_paragraph(BODY_TEXT)
    path = tmp_path / "toc.docx"
    d.save(str(path))
    tree = extract_document(str(path), get_settings()).section_tree
    ch1 = [n for n in tree.nodes if n.numbering == "1"]
    assert len(ch1) == 1  # the real heading, not the TOC listing


def test_image_only_docx_is_a_state_not_an_error(tmp_path):
    d = docx_lib.Document()
    d.add_picture(io.BytesIO(_png_bytes()))
    path = tmp_path / "scan.docx"
    d.save(str(path))
    result = extract_document(str(path), get_settings())
    assert result.image_only is True
    assert len(result.images) == 1


def test_malformed_docx_raises_taxonomy_error(tmp_path):
    bad = tmp_path / "broken.docx"
    bad.write_bytes(b"this is not a docx")
    with pytest.raises(FileMalformedError):
        extract_document(str(bad), get_settings())


def test_legacy_doc_rejected_with_clear_message():
    with pytest.raises(FileMalformedError) as exc_info:
        select_extractor(".doc")
    assert "'.doc'" in str(exc_info.value)
    assert ".docx" in str(exc_info.value)  # tells the user what IS supported
