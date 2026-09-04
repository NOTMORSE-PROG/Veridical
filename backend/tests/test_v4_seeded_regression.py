"""V4's own exit-checkpoint regression suite (TESTING.md §6 mechanism 4,
BOARD.md's V4 exit checkpoint): "a seeded-errors manuscript (fake DOI,
retracted paper, GRIM-impossible mean, wrong p) gets every seed flagged
with honest wording... a clean control manuscript stays clean."

This is the PERMANENT ship-blocker suite TESTING.md already names but
that didn't exist as a real file before this ticket — built here for the
first time. Runs the REAL pipeline (real ingestion, real citation/
forensics checks) against two DOCX manuscripts built in this file: one
with all four seeded errors, one with clean, verifiably-correct versions
of the same four claims. DOCX, not PDF -- kept that way even after
BUG-163 wired `page.find_tables()` into `ingest/pdf.py` (PDF native
table extraction is no longer the disclosed gap it was; see
`test_ingest_pdf.py::test_native_table_extracted_with_page_anchor_and_rows`
for that ingestion-level regression guard, and `tests/test_checks_forensics_extract.py`
for `extract_descriptive_stats`'s own format-agnostic coverage) -- DOCX's
`doc.add_table()` API stays the simpler, more deterministic way to build
THIS suite's specific fixture, not a statement about which format the
product actually supports.

Hits the real CrossRef API (DOI existence/retraction) — same "verify
against the real world, not just fixtures" standard V-027/V-028/V-029
were built and validated against. Gated on `DATABASE_URL` like every
other live-DB test in this suite; CrossRef itself needs no key.
"""

import os
from pathlib import Path

import docx as docx_lib
import pytest
from sqlalchemy import select, text
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine

from app.config import get_settings
from app.db import sqlalchemy_url
from app.llm.fake import FakeLLMClient
from app.models.enums import FlagSeverity, IngestStatus
from app.models.instructor import Instructor
from app.models.manuscript import Manuscript
from app.models.rubric import Rubric
from app.models.run import CheckResult, CheckRun, Flag
from app.pipeline.machine import run_check_run

live = pytest.mark.skipif(
    "DATABASE_URL" not in os.environ,
    reason="integration: needs a live Postgres + real internet (CrossRef) access",
)
pytestmark = live

SCRATCH_DB = "veridical_v4regressiontest"

# A real, confirmed-retracted paper (verified live 2026-08-06 against
# api.crossref.org/v1/works?filter=update-type:retraction — see V-028's
# ticket evidence). Using the REAL DOI is the point: proves the check
# calls the real CrossRef API and reads its real retraction metadata,
# not a mocked response.
_RETRACTED_DOI = "10.1016/j.micpro.2020.103768"
_RETRACTED_TITLE = (
    "Cross-Cultural communication of language learning social software "
    "based on FPGA and transfer learning"
)
# A real, confirmed-NOT-retracted paper (verified live against S2/CrossRef
# in V-028/V-030's own ticket evidence) for the clean control.
_CLEAN_DOI = "10.1038/nature12373"
_CLEAN_TITLE = "Nanometer scale thermometry in a living cell"

_FAKE_DOI = "10.9999/this-doi-does-not-exist-veridical-v4-seed"


def _seeded_docx(tmp_path) -> str:
    doc = docx_lib.Document()
    doc.add_heading("CHAPTER 4 RESULTS AND DISCUSSION", level=1)
    doc.add_paragraph(
        f"Prior work by Reyes (2026) established a baseline for capstone "
        f"quality assessment. A separate study, {'Zhang'} (2020), examined "
        f"related methods."
    )
    doc.add_paragraph(
        "The difference between groups was statistically significant, t(28) = 2.45, p = .500."
    )
    doc.add_paragraph("Table 1 summarizes the descriptive statistics for the two groups.")
    table = doc.add_table(rows=2, cols=3)
    for col, header in enumerate(("Group", "n", "M")):
        table.rows[0].cells[col].text = header
    for col, value in enumerate(("Respondents", "10", "3.33")):
        table.rows[1].cells[col].text = value

    doc.add_heading("REFERENCES", level=1)
    doc.add_paragraph(
        f"Reyes, J. P. (2026). A fabricated study on capstone quality. "
        f"Philippine Journal of Education, 15(2), 45-60. "
        f"https://doi.org/{_FAKE_DOI}"
    )
    doc.add_paragraph(
        f"Zhang, W. (2020). {_RETRACTED_TITLE}. Microprocessors and "
        f"Microsystems, 80, 103768. https://doi.org/{_RETRACTED_DOI}"
    )
    path = str(tmp_path / "seeded.docx")
    doc.save(path)
    return path


def _clean_control_docx(tmp_path) -> str:
    doc = docx_lib.Document()
    doc.add_heading("CHAPTER 4 RESULTS AND DISCUSSION", level=1)
    doc.add_paragraph(
        f"Prior work by Reyes (2026) established a baseline for capstone "
        f"quality assessment. A separate study, {'Kucsko'} (2013), examined "
        f"related methods."
    )
    doc.add_paragraph(
        "The difference between groups was statistically significant, t(28) = 2.45, p = .021."
    )
    doc.add_paragraph("Table 1 summarizes the descriptive statistics for the two groups.")
    table = doc.add_table(rows=2, cols=3)
    for col, header in enumerate(("Group", "n", "M")):
        table.rows[0].cells[col].text = header
    for col, value in enumerate(("Respondents", "10", "3.30")):
        table.rows[1].cells[col].text = value

    doc.add_heading("REFERENCES", level=1)
    # A real, unremarkable, non-retracted paper — DOI genuinely resolves.
    doc.add_paragraph(
        "Reyes, J. P. (2026). A genuine study on capstone quality. "
        "Philippine Journal of Education, 15(2), 45-60. "
        f"https://doi.org/{_CLEAN_DOI}"
    )
    doc.add_paragraph(
        f"Kucsko, G. (2013). {_CLEAN_TITLE}. Nature, 500, 54-58. https://doi.org/{_CLEAN_DOI}"
    )
    path = str(tmp_path / "clean.docx")
    doc.save(path)
    return path


@pytest.fixture(scope="module")
def scratch_url():
    import asyncio

    from alembic import command
    from tests.test_schema import _admin_execute, _alembic_config, _swap_db

    base = os.environ["DATABASE_URL"]
    asyncio.run(_admin_execute(base, f'DROP DATABASE IF EXISTS "{SCRATCH_DB}"'))
    asyncio.run(_admin_execute(base, f'CREATE DATABASE "{SCRATCH_DB}"'))
    url = _swap_db(base, SCRATCH_DB)
    command.upgrade(_alembic_config(url), "head")
    yield url
    asyncio.run(_admin_execute(base, f'DROP DATABASE IF EXISTS "{SCRATCH_DB}" WITH (FORCE)'))


@pytest.fixture()
def session_factory(scratch_url):
    engine = create_async_engine(sqlalchemy_url(scratch_url))
    yield async_sessionmaker(engine, expire_on_commit=False)


@pytest.fixture(autouse=True)
async def _clean_tables(session_factory):
    # Module-scoped scratch DB (module-scoped `scratch_url`, ticket AC:
    # one real CrossRef-hitting DB per module, not per test) — without
    # this, every test after the first collides on the seed instructor's
    # unique email (same pattern `test_flags_live.py` already uses).
    async with session_factory() as session:
        await session.execute(
            text(
                "TRUNCATE audit_log, flag, readiness_report, check_result, check_run, "
                "citation, criterion, rubric, manuscript, instructor RESTART IDENTITY CASCADE"
            )
        )
        await session.commit()
    yield


async def _seed_and_run(session_factory, tmp_path, monkeypatch, docx_path_fn):
    from app.ingest.service import ingest_manuscript

    monkeypatch.setenv("DATA_DIR", str(tmp_path / "data"))
    get_settings.cache_clear()
    settings = get_settings()
    docx_path = docx_path_fn(tmp_path)

    async with session_factory() as session:
        instructor = Instructor(email="v4-seed@test.local", display_name="V4 Seed Test")
        session.add(instructor)
        await session.commit()

        manuscript = Manuscript(
            instructor_id=instructor.id, group_label="Seed Group", file_ref=docx_path
        )
        session.add(manuscript)
        await session.commit()
        await ingest_manuscript(session, manuscript, Path(docx_path), settings)
        assert manuscript.ingest_status == IngestStatus.done

        rubric = Rubric(instructor_id=instructor.id, title="Format", source_file="r.docx")
        session.add(rubric)
        await session.commit()

        check_run = CheckRun(manuscript_id=manuscript.id, rubric_id=rubric.id)
        session.add(check_run)
        await session.commit()
        check_run_id = check_run.id

    async with session_factory() as session:
        check_run = await session.get(CheckRun, check_run_id)
        await run_check_run(session, check_run, settings, FakeLLMClient())

    async with session_factory() as verify:
        result_query = select(CheckResult).where(CheckResult.check_run_id == check_run_id)
        results = (await verify.execute(result_query)).scalars().all()
        flag_query = select(Flag).join(CheckResult).where(CheckResult.check_run_id == check_run_id)
        flags = (await verify.execute(flag_query)).scalars().all()
    return results, flags


async def test_seeded_manuscript_flags_all_four_planted_errors(
    session_factory, tmp_path, monkeypatch
):
    """V4 exit checkpoint: every one of the four seeded errors (fake DOI,
    retracted paper, GRIM-impossible mean, wrong p) must be flagged with
    honest wording and the arithmetic/evidence shown."""
    _, flags = await _seed_and_run(session_factory, tmp_path, monkeypatch, _seeded_docx)
    kinds = {f.detail.get("kind") for f in flags if f.detail}

    # Fake DOI -> unverifiable, never asserted as "fake" (charter rule 3).
    assert "unverifiable_not_found" in kinds
    # Real retracted paper -> retraction, the only high-severity F5 outcome.
    assert "retracted_source" in kinds
    # n=10, M=3.33 is not a GRIM-consistent value.
    assert "grim_inconsistent" in kinds
    # t(28)=2.45 recomputes to p=.021, not the reported p=.500.
    assert "p_value_mismatch" in kinds or "p_value_decision_error" in kinds
    assert len(flags) >= 4  # at minimum: fake DOI, retraction, GRIM, p-mismatch


async def test_clean_control_manuscript_produces_zero_false_positives(
    session_factory, tmp_path, monkeypatch
):
    """V4 exit checkpoint's other half — precision, not just recall: a
    manuscript with genuinely correct citations/statistics must get ZERO
    flags (charter judgment #1: a false accusation costs more than ten
    missed flags)."""
    _, flags = await _seed_and_run(session_factory, tmp_path, monkeypatch, _clean_control_docx)
    assert flags == []


async def test_retraction_flag_is_high_severity_with_honest_wording(
    session_factory, tmp_path, monkeypatch
):
    _, flags = await _seed_and_run(session_factory, tmp_path, monkeypatch, _seeded_docx)
    retraction_flags = [f for f in flags if f.detail and f.detail.get("kind") == "retracted_source"]
    assert len(retraction_flags) >= 1
    assert retraction_flags[0].severity == FlagSeverity.high
    reason = retraction_flags[0].detail["reason"]
    assert "RETRACTED" in reason
    # Honest wording (charter rule 3): a possible-inconsistency description,
    # never a direct accusation.
    assert "fabricat" not in reason.lower()
    assert "fake" not in reason.lower()


async def test_grim_impossible_mean_flagged_with_arithmetic(session_factory, tmp_path, monkeypatch):
    _, flags = await _seed_and_run(session_factory, tmp_path, monkeypatch, _seeded_docx)
    grim_flags = [f for f in flags if f.detail and f.detail.get("kind") == "grim_inconsistent"]
    assert len(grim_flags) >= 1
    assert "3.33" in grim_flags[0].evidence_excerpt


async def test_wrong_p_value_flagged(session_factory, tmp_path, monkeypatch):
    _, flags = await _seed_and_run(session_factory, tmp_path, monkeypatch, _seeded_docx)
    p_flags = [
        f
        for f in flags
        if f.detail and f.detail.get("kind") in ("p_value_mismatch", "p_value_decision_error")
    ]
    assert len(p_flags) >= 1
    assert "2.45" in p_flags[0].evidence_excerpt
